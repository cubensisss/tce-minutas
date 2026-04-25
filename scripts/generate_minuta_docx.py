"""
Script para gerar DOCX da minuta a partir do template Word da Conselheira.
Recebe os dados via JSON (stdin ou arquivo) e gera o documento final.
"""
import sys
import json
import argparse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'template.docx')

def clear_and_set(paragraph, text, bold=False):
    """Limpa o conteúdo do parágrafo e define novo texto."""
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        paragraph.runs[0].bold = bold
    else:
        run = paragraph.add_run(text)
        run.bold = bold

def add_paragraph_after(doc, reference_para, text, bold=False, italic=False, alignment=None, font_size=None):
    """Adiciona parágrafo após a referência."""
    new_para = doc.add_paragraph()
    reference_para._element.addnext(new_para._element)
    
    run = new_para.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    run.font.name = 'Arial'
    if alignment:
        new_para.alignment = alignment
    return new_para

def generate_docx(data, output_path):
    """Gera o DOCX a partir dos dados do processo."""
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Template não encontrado: {TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)
    p = doc.paragraphs
    
    processo = data.get('processo', {})
    ementa = data.get('ementa', '')
    relatorio = data.get('relatorio', '')
    analise = data.get('analise_completa', '')
    decisao = data.get('decisao_voto', '')
    
    # 1. Cabeçário / Identificação
    clear_and_set(p[2], f"PROCESSO TCE-PE Nº {processo.get('numero', '')}", bold=True)
    relator_raw = processo.get('relator', 'RODRIGO CAVALCANTI NOVAES').upper()
    relator_clean = re.sub(r'(CONSELHEIRO\s*)+', '', relator_raw).strip()
    clear_and_set(p[3], f"RELATOR: CONSELHEIRO {relator_clean}")
    clear_and_set(p[4], f"MODALIDADE - TIPO: Auditoria Especial - Conformidade - {processo.get('exercicio', '')}")
    clear_and_set(p[5], f"UNIDADE(S) JURISDICIONADA(S): {processo.get('unidade_jurisdicionada', '').upper()}")
    # Se p[6] era interessados ou um espaco, limpa se precisar, dependendo do template.
    # Normalmente, interessados fica algumas linhas abaixo após o cabecalho.
    clear_and_set(p[8], "INTERESSADOS:")
    clear_and_set(p[9], processo.get('interessados', ''))
    
    # 2. Ementa
    clear_and_set(p[11], ementa, bold=False)
    
    # 3. Descrição do Objeto (opcional)
    if processo.get('descricao_objeto'):
        clear_and_set(p[15], processo['descricao_objeto'])

    # 4. Truncar Documento de forma absoluta
    # O template original tem tabelas escondidas no final. Deletamos todos os elementos apos p[15].
    body = doc._body._element
    p15 = p[15]._element
    p15_index = body.index(p15)
    
    for child in list(body)[p15_index+1:]:
        body.remove(child)
        
    def add_p(text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, font_size=12):
        new_para = doc.add_paragraph()
        run = new_para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        run.font.name = 'Arial'
        new_para.alignment = alignment
        return new_para

    def add_bordered_title(title_text):
        t = doc.add_table(rows=1, cols=1)
        try:
            t.style = 'Table Grid'
        except:
            pass
        p_cell = t.cell(0, 0).paragraphs[0]
        p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_cell.add_run(title_text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'

    # 5. Inserir Relatório
    if relatorio:
        doc.add_paragraph() # espaco
        add_bordered_title("RELATÓRIO")
        
        lines = relatorio.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            clean = re.sub(r'[#*]+', '', line).strip()
            add_p(clean)
            
        add_p("É o relatório.")

    # 6. Inserir Análise (VOTO)
    if analise:
        doc.add_paragraph() # espaco
        add_bordered_title("VOTO")
        
        lines = analise.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            is_h = line.startswith('###')
            clean = re.sub(r'[#*]+', '', line).strip()
            add_p(clean, bold=is_h)

    # 7. Inserir Decisão
    if decisao:
        doc.add_paragraph() # espaco
        add_p("Ante o exposto, profiro o seguinte VOTO:", bold=True)
        
        lines = decisao.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            clean = re.sub(r'[#*]+', '', line).strip()
            is_bold = bool('CONSIDERANDO' in clean.upper() or re.match(r'^[IVX]+\.', clean))
            add_p(clean, bold=is_bold)

    doc.save(output_path)
    print(f"DOCX gerado: {output_path}")

if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--input', required=True)
    parser.add_argument('--output', required=True)
    args = parser.parse_args()
    
    with open(args.input, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    generate_docx(data, args.output)
